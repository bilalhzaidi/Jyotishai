"""Report generation utilities for JyotishAI.

This module assembles analysis results into human‑readable reports and
exports them as DOCX or PDF files.  The generated files are saved in
the `output/` directory.
"""

from __future__ import annotations

import os
import datetime
from typing import List

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
except ImportError:
    Document = None  # type: ignore
    Inches = None  # type: ignore

try:
    from reportlab.lib.pagesizes import LETTER  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
except ImportError:
    LETTER = None  # type: ignore
    getSampleStyleSheet = None  # type: ignore
    inch = None  # type: ignore
    canvas = None  # type: ignore
    SimpleDocTemplate = None  # type: ignore
    Paragraph = None  # type: ignore
    Spacer = None  # type: ignore

from models.schemas import ModuleResult, AnalysisResponse
from services.chart_engine import VedicAstrologyEngine

# Directory to save generated reports
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _timestamp_slug(name: str) -> str:
    ts = datetime.datetime.utcnow().strftime("%Y%m%d%H%M%S")
    slug = name.lower().replace(" ", "_")
    return f"{slug}_{ts}"


def generate_docx(chart: VedicAstrologyEngine, results: List[ModuleResult]) -> str:
    """Generate a DOCX report and return the file path."""
    filename = _timestamp_slug(chart.name) + ".docx"
    path = os.path.join(OUTPUT_DIR, filename)
    if Document is None:
        # Fallback: create a simple text file with .docx extension
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"Astrological Report for {chart.name}\n")
            f.write(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(
                f"Birth Details: {chart.birth_datetime.strftime('%Y-%m-%d %H:%M:%S')} (UTC{chart.utc_offset:+.1f}), "
                f"Location: {chart.location}\n\n"
            )
            for result in results:
                f.write(f"{result.module.value}\n")
                f.write(result.analysis + "\n\n")
        return path
    # Use python-docx if available
    doc = Document()
    doc.add_heading(f"Astrological Report for {chart.name}", level=1)
    doc.add_paragraph(
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    doc.add_paragraph(
        f"Birth Details: {chart.birth_datetime.strftime('%Y-%m-%d %H:%M:%S')} (UTC{chart.utc_offset:+.1f}), "
        f"Location: {chart.location}"
    )
    doc.add_paragraph("")
    for result in results:
        doc.add_heading(result.module.value, level=2)
        doc.add_paragraph(result.analysis)
        doc.add_paragraph("")
    doc.save(path)
    return path


def generate_pdf(chart: VedicAstrologyEngine, results: List[ModuleResult]) -> str:
    """Generate a PDF report and return the file path."""
    filename = _timestamp_slug(chart.name) + ".pdf"
    path = os.path.join(OUTPUT_DIR, filename)
    if SimpleDocTemplate is None or getSampleStyleSheet is None:
        # Fallback: generate a plain text file with .pdf extension
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"Astrological Report for {chart.name}\n")
            f.write(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(
                f"Birth Details: {chart.birth_datetime.strftime('%Y-%m-%d %H:%M:%S')} (UTC{chart.utc_offset:+.1f}), "
                f"Location: {chart.location}\n\n"
            )
            for result in results:
                f.write(f"{result.module.value}\n")
                f.write(result.analysis + "\n\n")
        return path
    # Build a simple report using reportlab's platypus
    doc = SimpleDocTemplate(path, pagesize=LETTER, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []
    title_style = styles['Title']
    normal_style = styles['BodyText']
    heading_style = styles['Heading2']

    story.append(Paragraph(f"Astrological Report for {chart.name}", title_style))
    story.append(Spacer(1, 0.2 * inch))
    story.append(
        Paragraph(
            f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            normal_style,
        )
    )
    story.append(
        Paragraph(
            f"Birth Details: {chart.birth_datetime.strftime('%Y-%m-%d %H:%M:%S')} (UTC{chart.utc_offset:+.1f}), "
            f"Location: {chart.location}",
            normal_style,
        )
    )
    story.append(Spacer(1, 0.2 * inch))
    for result in results:
        story.append(Paragraph(result.module.value, heading_style))
        story.append(Paragraph(result.analysis, normal_style))
        story.append(Spacer(1, 0.2 * inch))
    doc.build(story)
    return path